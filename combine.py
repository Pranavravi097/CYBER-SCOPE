import os
import socket
import time
import ssl
import requests
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Inches
from flask import Flask, request, jsonify, render_template_string
from flask_cors import CORS
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from urllib.parse import urljoin
import re

app = Flask(__name__)
CORS(app)

# HTML Template as a string
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Website Security Check Using Selenium</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            background-color: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        }
        h1 {
            color: #333;
            text-align: center;
        }
        h2 {
            color: #444;
            margin-top: 30px;
        }
        .url-input {
            width: 100%;
            max-width: 500px;
            padding: 10px;
            margin: 20px 0;
            border: 2px solid #ccc;
            border-radius: 4px;
            font-size: 16px;
        }
        .check-button {
            padding: 12px 24px;
            background-color: #007bff;
            color: white;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
            transition: background-color 0.3s;
        }
        .check-button:hover {
            background-color: #0056b3;
        }
        .result-box {
            margin-top: 20px;
            padding: 15px;
            border-radius: 4px;
            display: none;
        }
        .safe {
            background-color: #d4edda;
            border: 1px solid #c3e6cb;
            color: #155724;
        }
        .unsafe {
            background-color: #f8d7da;
            border: 1px solid #f5c6cb;
            color: #721c24;
        }
        .reasons-list {
            margin-top: 10px;
            padding-left: 20px;
        }
        .reasons-list li {
            margin-bottom: 5px;
        }
        .section {
            margin-bottom: 40px;
            padding: 20px;
            background-color: white;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        }
        #message {
            margin-top: 20px;
            padding: 10px;
            border-radius: 4px;
            text-align: center;
        }
        .success {
            background-color: #d4edda;
            color: #155724;
        }
        .error {
            background-color: #f8d7da;
            color: #721c24;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Website Security Scanner</h1>
        <div class="section">
            <h2>Check URL Safety</h2>
            <input type="text" id="urlInput" class="url-input" placeholder="Enter URL to check (e.g., https://example.com)">
            <button onclick="checkURL()" class="check-button">Check URL</button>
            <div id="resultBox" class="result-box">
                <div id="resultMessage"></div>
                <ul id="reasonsList" class="reasons-list"></ul>
            </div>
        </div>
        <div class="section">
            <h2>Generate Security Report</h2>
            <p>Click the button below to take screenshots and generate a comprehensive security report for the configured websites.</p>
            <button onclick="takeScreenshot()" class="check-button">Generate Report</button>
            <div id="message"></div>
        </div>
    </div>
    <script>
        async function checkURL() {
            const urlInput = document.getElementById('urlInput').value;
            const resultBox = document.getElementById('resultBox');
            const resultMessage = document.getElementById('resultMessage');
            const reasonsList = document.getElementById('reasonsList');
            
            try {
                const response = await fetch('/check_url', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({ url: urlInput })
                });
                
                const data = await response.json();
                
                // Clear previous results
                reasonsList.innerHTML = '';
                
                // Show result box
                resultBox.style.display = 'block';
                resultBox.className = 'result-box ' + (data.is_safe ? 'safe' : 'unsafe');
                
                // Set message
                resultMessage.textContent = data.is_safe ? 
                    '✓ This URL appears to be safe' : 
                    '⚠ This URL may not be safe';
                
                // Add reasons if any
                if (data.reasons && data.reasons.length > 0) {
                    data.reasons.forEach(reason => {
                        const li = document.createElement('li');
                        li.textContent = reason;
                        reasonsList.appendChild(li);
                    });
                }
                
                // Change input border color
                document.getElementById('urlInput').style.borderColor = data.is_safe ? '#28a745' : '#dc3545';
                
            } catch (error) {
                resultBox.style.display = 'block';
                resultBox.className = 'result-box unsafe';
                resultMessage.textContent = 'Error checking URL';
                const li = document.createElement('li');
                li.textContent = error.message;
                reasonsList.appendChild(li);
            }
        }

        async function takeScreenshot() {
            const messageDiv = document.getElementById('message');
            messageDiv.textContent = 'Processing...';
            messageDiv.className = '';
            
            try {
                const response = await fetch('/take_screenshot');
                const data = await response.json();
                messageDiv.textContent = data.message;
                messageDiv.className = data.status === 'success' ? 'success' : 'error';
            } catch (error) {
                messageDiv.textContent = 'Error: ' + error.message;
                messageDiv.className = 'error';
            }
        }
    </script>
</body>
</html>
"""

# Configuration
urls = [
    "https://www.blackbox.ai",
    "https://example.com",
    "https://www.w3schools.com",
    "https://www.python.org",
    "https://www.selenium.dev"
]

chromedriver_path = r"C:\Users\MARTTEN\Desktop\final yr project\project_env\chromedriver-win64\chromedriver.exe"
output_dir = r"C:\Users\MARTTEN\Desktop\final yr project\Screenshots"

def is_url_safe(url):
    """
    Check if a URL is safe by performing various security checks.
    Returns a tuple (is_safe, reasons)
    """
    try:
        # Validate URL format
        if not re.match(r'https?://[\w\-\.]+\.[a-zA-Z]{2,}/?.*', url):
            return False, ["Invalid URL format"]

        # Initialize reasons list for detailed feedback
        safety_issues = []
        
        # Make request with a timeout
        response = requests.get(url, verify=False, timeout=10)
        
        # Check HTTPS
        if not url.startswith('https://'):
            safety_issues.append("No HTTPS connection")
            
        # Check security headers
        headers = response.headers
        if 'Strict-Transport-Security' not in headers:
            safety_issues.append("Missing HSTS header")
        if 'X-Content-Type-Options' not in headers:
            safety_issues.append("Missing security headers")
            
        # Check SSL/TLS
        try:
            hostname = url.split("://")[1].split("/")[0]
            context = ssl.create_default_context()
            with context.wrap_socket(socket.socket(), server_hostname=hostname) as s:
                s.connect((hostname, 443))
                cert = s.getpeercert()
                
                # Check certificate expiration
                cert_expiry = ssl.cert_time_to_seconds(cert['notAfter'])
                if time.time() > cert_expiry:
                    safety_issues.append("SSL Certificate expired")
        except Exception as e:
            safety_issues.append(f"SSL/TLS verification failed: {str(e)}")
            
        # Consider the URL safe if there are no major security issues
        is_safe = len(safety_issues) <= 1  # Allow one minor issue
        
        return is_safe, safety_issues
        
    except requests.exceptions.RequestException as e:
        return False, [f"Connection error: {str(e)}"]
    except Exception as e:
        return False, [f"Error checking URL: {str(e)}"]

def check_vulnerabilities(url):
    vulnerabilities = []
    
    try:
        # 1. Check HTTPS
        response = requests.get(url, verify=False, timeout=10)
        if not url.startswith('https://'):
            vulnerabilities.append("No HTTPS: Site doesn't use secure connection")
            
        # 2. Check Security Headers
        headers = response.headers
        security_headers = {
            'Strict-Transport-Security': 'Missing HSTS header',
            'X-Content-Type-Options': 'Missing X-Content-Type-Options header',
            'X-Frame-Options': 'Missing X-Frame-Options header',
            'Content-Security-Policy': 'Missing Content Security Policy'
        }
        
        for header, message in security_headers.items():
            if header not in headers:
                vulnerabilities.append(message)
                
        # 3. Check SSL/TLS Configuration
        try:
            hostname = url.split("://")[1].split("/")[0]
            context = ssl.create_default_context()
            with context.wrap_socket(socket.socket(), server_hostname=hostname) as s:
                s.connect((hostname, 443))
                cert = s.getpeercert()
                
                # Check certificate expiration
                cert_expiry = ssl.cert_time_to_seconds(cert['notAfter'])
                if time.time() > cert_expiry:
                    vulnerabilities.append("SSL Certificate has expired")
                
        except Exception as e:
            vulnerabilities.append(f"SSL/TLS Error: {str(e)}")
            
        # 4. Check for Common Sensitive Files
        sensitive_files = ['/robots.txt', '/sitemap.xml', '/.env', '/.git/config']
        for file in sensitive_files:
            try:
                file_url = urljoin(url, file)
                response = requests.get(file_url, timeout=5)
                if response.status_code == 200:
                    vulnerabilities.append(f"Sensitive file exposed: {file}")
            except:
                pass
                
    except Exception as e:
        vulnerabilities.append(f"Error during vulnerability check: {str(e)}")
        
    return vulnerabilities

def check_risk_level(url):
    """
    Assess the risk level of a URL.
    Returns a tuple (risk_level, reasons)
    """
    risk_level = "Low"
    reasons = []

    # Example criteria for risk assessment
    if not url.startswith('https://'):
        risk_level = "Medium"
        reasons.append("No HTTPS connection")

    # Check for common vulnerabilities
    vulnerabilities = check_vulnerabilities(url)
    if vulnerabilities:
        risk_level = "High"
        reasons.append("Vulnerabilities detected: " + ", ".join(vulnerabilities))

    return risk_level, reasons

def send_email(subject, body, to_email, attachment_path=None):
    from_email = "pranav09072004@gmail.com"
    password = "xome witg ergv rftp"  # Consider using environment variables for sensitive data

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    if attachment_path:
        try:
            with open(attachment_path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())
            
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {attachment_path.split('/')[-1]}",
            )
            msg.attach(part)
            print(f"Attachment {attachment_path} added.")
            
        except Exception as e:
            print(f"Failed to attach file: {e}")

    server = None
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        server.send_message(msg)
        print("Email sent successfully!")
        
    except Exception as e:
        print(f"Failed to send email: {e}")
        
    finally:
        if server:
            server.quit()

@app.route('/')
def home():
    return render_template_string(HTML_TEMPLATE)

@app.route('/check_url', methods=['POST'])
def check_url():
    """
    Endpoint to check if a URL is safe
    """
    data = request.get_json()
    url = data.get('url', '')
    
    if not url:
        return jsonify({
            "status": "error",
            "message": "No URL provided",
            "is_safe": False
        })
    
    is_safe, reasons = is_url_safe(url)
    
    return jsonify({
        "status": "success",
        "is_safe": is_safe,
        "reasons": reasons,
        "color": "green" if is_safe else "red"
    })

@app.route('/take_screenshot', methods=['GET', 'POST'])
def take_screenshot():
    try:
        chrome_options = Options()
        chrome_options.add_argument("--start-maximized")
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--disable-extensions")
        chrome_options.add_argument("--disable-gpu")
        
        os.makedirs(output_dir, exist_ok=True)
        
        service = Service(chromedriver_path)
        driver = webdriver.Chrome(service=service, options=chrome_options)
        
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        document_path = "website_report.docx"
        renamed_document_path = f"website_report_{timestamp}.docx"
        
        if os.path.exists(document_path):
            os.rename(document_path, renamed_document_path)
            print(f"Renamed existing document to {renamed_document_path}")
        
        doc = Document()
        doc.add_heading("Website Security and Screenshot Report", level=1)
        
        for url in urls:
            print(f"Processing {url}...")
            
            # Add URL heading
            doc.add_heading(f"Website: {url}", level=2)
            
            # Take screenshot
            driver.get(url)
            time.sleep(5)
            screenshot_path = os.path.join(output_dir, f"screenshot_{timestamp}_{url.replace('://', '_').replace('/', '_').replace('.', '-')}.png")
            driver.save_screenshot(screenshot_path)
            
            # Add screenshot to document
            if os.path.exists(screenshot_path):
                doc.add_picture(screenshot_path, width=Inches(6))
            
            # Check vulnerabilities
            vulnerabilities = check_vulnerabilities(url)

            # Check risk level
            risk_level, risk_reasons = check_risk_level(url)

            # Add vulnerability information
            doc.add_heading("Security Analysis:", level=3)
            if vulnerabilities:
                doc.add_paragraph("The following vulnerabilities were detected:")
                for vuln in vulnerabilities:
                    doc.add_paragraph(vuln, style='List Bullet')
            else:
                doc.add_paragraph("No major vulnerabilities detected.")

            # Add risk level information
            doc.add_heading("Risk Level Assessment:", level=3)
            doc.add_paragraph(f"Risk Level: {risk_level}")
            if risk_reasons:
                doc.add_paragraph("Reasons for risk level assessment:")
                for reason in risk_reasons:
                    doc.add_paragraph(reason, style='List Bullet')

            doc.add_paragraph("\n")  # Add spacing between websites
        
        # Save the document
        doc.save(document_path)
        print(f"Report saved as {document_path}")
        
        # Send email with report
        to_email = "pranav09072004@gmail.com"  # Replace with actual email
        send_email(
            subject="Website Security Report",
            body="Please find attached the website security report.",
            to_email=to_email,
            attachment_path=document_path
        )
        
        driver.quit()
        
        return jsonify({
            "status": "success",
            "message": "Screenshots taken and report generated successfully!"
        })
        
    except Exception as e:
        if 'driver' in locals():
            driver.quit()
        return jsonify({
            "status": "error",
            "message": f"Error: {str(e)}"
        })

if __name__ == '__main__':
    app.run(debug=True)